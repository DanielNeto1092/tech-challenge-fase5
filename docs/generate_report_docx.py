"""Gera o relatório técnico aprofundado e editável da Guardiã AI em DOCX."""

from __future__ import annotations

import importlib
import json
import math
import subprocess
import sys
import tempfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from textwrap import shorten
from typing import Any
from zoneinfo import ZoneInfo

import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
DATASET = BACKEND / "data" / "raw" / "maternal_health_risk.csv"
TRAINING_REPORT = BACKEND / "artifacts" / "training_report_v1.0.0.json"
MODEL_ARTIFACT = BACKEND / "artifacts" / "maternal_risk_model_v1.0.0.joblib"
KNOWLEDGE_BASE = BACKEND / "data" / "knowledge_base"
OUTPUT = ROOT / "docs" / "relatorio-tecnico.docx"

NAVY = "183B56"
TEAL = "148A85"
PALE_TEAL = "E8F5F3"
INK = "24333D"
MUTED = "5E6D75"
LIGHT = "EFF3F4"
RED = "A13D4A"
AMBER = "B07318"
WHITE = "FFFFFF"

FEATURE_COLUMNS = (
    "age",
    "systolic_bp",
    "diastolic_bp",
    "blood_sugar",
    "body_temperature",
    "heart_rate",
)
FEATURE_LABELS = {
    "age": "Idade",
    "systolic_bp": "Pressão sistólica",
    "diastolic_bp": "Pressão diastólica",
    "blood_sugar": "Glicemia",
    "body_temperature": "Temperatura corporal",
    "heart_rate": "Frequência cardíaca",
}
FEATURE_UNITS = {
    "age": "anos",
    "systolic_bp": "mmHg",
    "diastolic_bp": "mmHg",
    "blood_sugar": "mmol/L",
    "body_temperature": "°F",
    "heart_rate": "bpm",
}
RISK_LABELS = {0: "Baixo", 1: "Médio", 2: "Alto"}
MODEL_LABELS = {
    "multinomial_logistic_regression": "Regressão logística multinomial",
    "random_forest": "Random forest",
}
MONTHS = {
    1: "janeiro",
    2: "fevereiro",
    3: "março",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro",
}


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_cell_text(cell, value: Any, *, bold=False, color=INK, size=8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[Any]],
    *,
    font_size: float = 8.5,
    widths: list[float] | None = None,
) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    _repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        _shade(table.rows[0].cells[index], NAVY)
        _set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
            color=WHITE,
            size=font_size,
        )
        if widths:
            table.rows[0].cells[index].width = Cm(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            if row_index % 2:
                _shade(cells[column_index], LIGHT)
            _set_cell_text(cells[column_index], value, size=font_size)
            if widths:
                cells[column_index].width = Cm(widths[column_index])
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _add_heading(document: Document, text: str, level: int = 1) -> Any:
    paragraph = document.add_heading(text, level=level)
    _set_keep_with_next(paragraph)
    return paragraph


def _add_body(
    document: Document,
    text: str,
    *,
    bold_prefix: str | None = None,
    italic=False,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        run = paragraph.add_run(text[len(bold_prefix) :])
    else:
        run = paragraph.add_run(text)
    run.italic = italic
    return paragraph


def _add_bullet(document: Document, text: str, *, level=0) -> Any:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(text, style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def _add_number(document: Document, text: str) -> Any:
    paragraph = document.add_paragraph(text, style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def _add_callout(document: Document, text: str, *, tone="teal") -> Any:
    palette = {
        "teal": (PALE_TEAL, TEAL),
        "warning": ("FFF5E6", AMBER),
        "danger": ("FBECEF", RED),
    }
    fill, border = palette[tone]
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, fill)
    properties = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "10")
        edge.set(qn("w:color"), border)
        borders.append(edge)
    properties.append(borders)
    _set_cell_text(cell, text, bold=True, color=NAVY, size=9.3)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.runs[0]
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def _add_picture(document: Document, path: Path, caption: str, *, width=6.45) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    _add_caption(document, caption)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def _add_toc(document: Document) -> None:
    _add_heading(document, "Sumário", 1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Atualize este campo no Word para exibir o sumário e as páginas."
    separate.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    document.add_page_break()


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)

    heading_colors = {1: NAVY, 2: TEAL, 3: MUTED}
    heading_sizes = {1: 17, 2: 13, 3: 10.5}
    for level in (1, 2, 3):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.size = Pt(heading_sizes[level])
        style.font.color.rgb = RGBColor.from_string(heading_colors[level])
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(9.2)

    header = section.header.paragraphs[0]
    header.text = "GUARDIÃ AI  |  Relatório técnico aprofundado"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    _add_page_number(section.footer.paragraphs[0])

    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _percent(value: float, digits=2) -> str:
    return f"{value * 100:.{digits}f}%".replace(".", ",")


def _number(value: float, digits=2) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def _wilson_interval(successes: int, total: int, z=1.96) -> tuple[float, float]:
    if total <= 0:
        return 0.0, 0.0
    proportion = successes / total
    denominator = 1 + (z**2 / total)
    center = (proportion + z**2 / (2 * total)) / denominator
    margin = (
        z * math.sqrt((proportion * (1 - proportion) / total) + z**2 / (4 * total**2)) / denominator
    )
    return center - margin, center + margin


def _per_class_metrics(matrix: list[list[int]]) -> list[dict[str, float | int | str]]:
    values = np.asarray(matrix, dtype=float)
    rows = []
    for class_id in range(3):
        true_positive = values[class_id, class_id]
        support = values[class_id, :].sum()
        predicted = values[:, class_id].sum()
        precision = true_positive / predicted if predicted else 0.0
        recall = true_positive / support if support else 0.0
        f1 = 2 * precision * recall / (precision + recall) if precision + recall else 0.0
        rows.append(
            {
                "class": RISK_LABELS[class_id],
                "precision": float(precision),
                "recall": float(recall),
                "f1": float(f1),
                "support": int(support),
            }
        )
    return rows


def _git_output(*args: str) -> str:
    try:
        return subprocess.check_output(
            ["git", *args],
            cwd=ROOT,
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
    except (OSError, subprocess.CalledProcessError):
        return "não disponível"


def _is_tracked(path: str) -> bool:
    try:
        subprocess.run(
            ["git", "ls-files", "--error-unmatch", path],
            cwd=ROOT,
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return True
    except (OSError, subprocess.CalledProcessError):
        return False


def _load_context() -> dict[str, Any]:
    report = json.loads(TRAINING_REPORT.read_text(encoding="utf-8"))
    raw = pd.read_csv(DATASET).rename(
        columns={
            "Age": "age",
            "SystolicBP": "systolic_bp",
            "DiastolicBP": "diastolic_bp",
            "Blood glucose": "blood_sugar",
            "BodyTemp": "body_temperature",
            "HeartRate": "heart_rate",
            "RiskLevel": "risk_level",
        }
    )
    clean = raw.drop_duplicates(keep="first").reset_index(drop=True)

    grouped = clean.groupby(list(FEATURE_COLUMNS), dropna=False)["risk_level"]
    label_sets = grouped.apply(lambda values: tuple(sorted({int(item) for item in values})))
    conflicting_keys = label_sets[label_sets.apply(len) > 1]
    conflict_combinations = Counter(
        "–".join(RISK_LABELS[item] for item in labels) for labels in conflicting_keys
    )
    conflict_rows = int(
        clean.set_index(list(FEATURE_COLUMNS)).index.isin(conflicting_keys.index).sum()
    )

    section_counts: dict[str, int] = {}
    source_metadata: list[dict[str, Any]] = []
    for path in sorted(KNOWLEDGE_BASE.glob("*.json")):
        payload = json.loads(path.read_text(encoding="utf-8"))
        section_counts[payload["source_id"]] = len(payload["sections"])
        source_metadata.append(
            {
                "title": payload["title"],
                "issuer": payload.get("issuer", "Não informado"),
                "year": payload.get("year", "Não informado"),
                "license": payload.get("license", "Não explicitada"),
                "sections": len(payload["sections"]),
                "url": payload["url"],
            }
        )

    sys.path.insert(0, str(BACKEND))
    predictor_module = importlib.import_module("app.ml")
    retriever_module = importlib.import_module("app.rag.retriever")
    predictor = predictor_module.MaternalRiskPredictor.load(MODEL_ARTIFACT)
    example_features = {
        "age": 35,
        "systolic_bp": 140,
        "diastolic_bp": 90,
        "blood_sugar": 13.0,
        "body_temperature": 98.0,
        "heart_rate": 70,
    }
    prediction = predictor.predict(example_features)
    explanation = predictor.explain(
        example_features,
        target_class=int(prediction["risk_level"]),
    )
    retriever = retriever_module.ProtocolRetriever(KNOWLEDGE_BASE)
    retrieval_query = (
        "O que os protocolos informam sobre a pressão arterial? risco materno Alto "
        "gestação Pressão arterial sistólica, Glicemia, Idade"
    )
    retrieval = retriever.retrieve(retrieval_query, top_k=4)

    return {
        "report": report,
        "raw": raw,
        "clean": clean,
        "conflict_combinations": conflict_combinations,
        "conflict_rows": conflict_rows,
        "source_metadata": source_metadata,
        "section_counts": section_counts,
        "example_features": example_features,
        "prediction": prediction,
        "explanation": explanation,
        "retrieval_query": retrieval_query,
        "retrieval": retrieval,
        "commit": _git_output("rev-parse", "--short", "HEAD"),
        "branch": _git_output("branch", "--show-current"),
        "readme_tracked": _is_tracked("README.md"),
    }


def _chart_style() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 9,
            "axes.titlesize": 11,
            "axes.labelsize": 9,
            "axes.edgecolor": "#B8C5CA",
            "axes.labelcolor": "#24333D",
            "xtick.color": "#45545C",
            "ytick.color": "#45545C",
            "figure.facecolor": "white",
            "axes.facecolor": "white",
        }
    )


def _save_figure(path: Path) -> None:
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()


def _create_charts(context: dict[str, Any], destination: Path) -> dict[str, Path]:
    _chart_style()
    paths: dict[str, Path] = {}
    raw = context["raw"]
    clean = context["clean"]
    report = context["report"]

    paths["classes"] = destination / "classes.png"
    raw_counts = raw["risk_level"].value_counts().reindex([0, 1, 2], fill_value=0)
    clean_counts = clean["risk_level"].value_counts().reindex([0, 1, 2], fill_value=0)
    x = np.arange(3)
    plt.figure(figsize=(7.1, 3.7))
    plt.bar(x - 0.19, raw_counts, 0.38, label="Arquivo bruto", color="#718EA1")
    plt.bar(x + 0.19, clean_counts, 0.38, label="Após deduplicação", color="#148A85")
    for position, value in zip(x - 0.19, raw_counts, strict=True):
        plt.text(position, value + 10, str(value), ha="center", fontsize=8)
    for position, value in zip(x + 0.19, clean_counts, strict=True):
        plt.text(position, value + 10, str(value), ha="center", fontsize=8)
    plt.xticks(x, ["Baixo", "Médio", "Alto"])
    plt.ylabel("Registros")
    plt.title("Distribuição das classes antes e depois da deduplicação")
    plt.legend(frameon=False)
    plt.grid(axis="y", alpha=0.18)
    _save_figure(paths["classes"])

    paths["distributions"] = destination / "distributions.png"
    figure, axes = plt.subplots(3, 2, figsize=(7.2, 8.0))
    for axis, feature in zip(axes.flat, FEATURE_COLUMNS, strict=True):
        axis.hist(clean[feature], bins=min(18, int(clean[feature].nunique())), color="#148A85")
        axis.set_title(f"{FEATURE_LABELS[feature]} ({FEATURE_UNITS[feature]})")
        axis.set_ylabel("Frequência")
        axis.grid(axis="y", alpha=0.16)
    figure.suptitle("Distribuições dos atributos após deduplicação", fontsize=13)
    figure.tight_layout(rect=(0, 0, 1, 0.97))
    figure.savefig(paths["distributions"], dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(figure)

    paths["correlation"] = destination / "correlation.png"
    correlation = clean[[*FEATURE_COLUMNS, "risk_level"]].corr(numeric_only=True)
    labels = [*[FEATURE_LABELS[item] for item in FEATURE_COLUMNS], "Classe"]
    figure, axis = plt.subplots(figsize=(7.2, 5.5))
    image = axis.imshow(correlation, cmap="RdBu_r", vmin=-1, vmax=1)
    axis.set_xticks(range(len(labels)), labels, rotation=40, ha="right")
    axis.set_yticks(range(len(labels)), labels)
    for row in range(len(labels)):
        for column in range(len(labels)):
            value = correlation.iloc[row, column]
            color = "white" if abs(value) > 0.55 else "#24333D"
            axis.text(column, row, f"{value:.2f}", ha="center", va="center", color=color)
    axis.set_title("Correlação linear exploratória")
    figure.colorbar(image, ax=axis, fraction=0.046, pad=0.04)
    _save_figure(paths["correlation"])

    paths["conflicts"] = destination / "conflicts.png"
    combinations = context["conflict_combinations"]
    plt.figure(figsize=(7.0, 3.5))
    names = list(combinations)
    values = [combinations[name] for name in names]
    bars = plt.barh(names, values, color=["#718EA1", "#A13D4A", "#B07318", "#6C5B7B"])
    for bar, value in zip(bars, values, strict=True):
        plt.text(value + 0.25, bar.get_y() + bar.get_height() / 2, str(value), va="center")
    plt.xlabel("Grupos de vetores idênticos")
    plt.title("Composição dos 35 grupos com rótulos conflitantes")
    plt.grid(axis="x", alpha=0.16)
    _save_figure(paths["conflicts"])

    paths["confusion"] = destination / "confusion.png"
    figure, axes = plt.subplots(1, 2, figsize=(7.4, 3.7))
    for axis, model_name in zip(axes, report["candidate_models"], strict=True):
        matrix = np.asarray(report["evaluations"][model_name]["confusion_matrix"])
        image = axis.imshow(matrix, cmap="Blues")
        axis.set_xticks(range(3), ["Baixo", "Médio", "Alto"])
        axis.set_yticks(range(3), ["Baixo", "Médio", "Alto"])
        axis.set_xlabel("Classe prevista")
        axis.set_ylabel("Classe real")
        axis.set_title(MODEL_LABELS[model_name])
        threshold = matrix.max() / 2
        for row in range(3):
            for column in range(3):
                axis.text(
                    column,
                    row,
                    str(matrix[row, column]),
                    ha="center",
                    va="center",
                    color="white" if matrix[row, column] > threshold else "#24333D",
                    fontweight="bold",
                )
        figure.colorbar(image, ax=axis, fraction=0.046, pad=0.04)
    figure.suptitle("Matrizes de confusão no holdout", fontsize=13)
    figure.tight_layout(rect=(0, 0, 1, 0.94))
    figure.savefig(paths["confusion"], dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(figure)

    paths["importance"] = destination / "importance.png"
    importance = report["global_feature_importance"]
    ordered = list(reversed(list(importance.items())))
    plt.figure(figsize=(7.0, 3.8))
    bars = plt.barh(
        [FEATURE_LABELS[name] for name, _ in ordered],
        [value for _, value in ordered],
        color="#148A85",
    )
    for bar, (_, value) in zip(bars, ordered, strict=True):
        plt.text(value + 0.005, bar.get_y() + bar.get_height() / 2, _percent(value), va="center")
    plt.xlabel("Importância de impureza normalizada")
    plt.title("Importância global do modelo selecionado")
    plt.xlim(0, max(importance.values()) * 1.22)
    plt.grid(axis="x", alpha=0.16)
    _save_figure(paths["importance"])

    paths["waterfall"] = destination / "waterfall.png"
    explanation = context["explanation"]
    contributions = explanation["feature_contributions"]
    ordered_contributions = sorted(
        contributions.items(),
        key=lambda item: abs(item[1]),
        reverse=True,
    )
    values = [contribution * 100 for _, contribution in ordered_contributions]
    labels_waterfall = [FEATURE_LABELS[feature] for feature, _ in ordered_contributions]
    colors = ["#148A85" if value >= 0 else "#A13D4A" for value in values]
    y = np.arange(len(labels_waterfall))
    figure, axis = plt.subplots(figsize=(7.2, 4.5))
    bars = axis.barh(y, values, color=colors)
    limit = max(abs(value) for value in values) * 1.30
    axis.set_xlim(-limit, limit)
    for bar, value in zip(bars, values, strict=True):
        offset = limit * 0.025
        axis.text(
            value + (offset if value >= 0 else -offset),
            bar.get_y() + bar.get_height() / 2,
            f"{value:+.2f} pp",
            ha="left" if value >= 0 else "right",
            va="center",
            color="#24333D",
            fontsize=8,
            fontweight="bold",
        )
    axis.axvline(0, color="#48606F", linewidth=1)
    axis.set_yticks(y, labels_waterfall)
    axis.invert_yaxis()
    axis.set_xlabel("Contribuição para a probabilidade da classe alta (pontos percentuais)")
    axis.set_title("Contribuições locais — estudo de caso")
    figure.text(
        0.5,
        0.01,
        (
            f"Baseline {_percent(explanation['baseline_probability'])} + contribuições "
            f"= predição {_percent(explanation['model_probability'])}"
        ),
        ha="center",
        color="#5E6D75",
        fontsize=8.5,
    )
    axis.grid(axis="x", alpha=0.16)
    figure.tight_layout(rect=(0, 0.05, 1, 1))
    figure.savefig(paths["waterfall"], dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(figure)

    paths["retrieval"] = destination / "retrieval.png"
    retrieval = context["retrieval"]
    source_labels = [
        f"[{index}] {shorten(item.reference, width=47, placeholder='…')}"
        for index, item in enumerate(retrieval, 1)
    ]
    scores = [item.relevance_score for item in retrieval]
    plt.figure(figsize=(7.2, 3.5))
    bars = plt.barh(list(reversed(source_labels)), list(reversed(scores)), color="#718EA1")
    for bar, score in zip(bars, reversed(scores), strict=True):
        plt.text(score + 0.003, bar.get_y() + bar.get_height() / 2, f"{score:.4f}", va="center")
    plt.xlabel("Similaridade de cosseno")
    plt.title("Recuperação documental para a pergunta do estudo de caso")
    plt.xlim(0, max(scores) * 1.25)
    plt.grid(axis="x", alpha=0.16)
    _save_figure(paths["retrieval"])

    paths["architecture"] = destination / "architecture.png"
    figure, axis = plt.subplots(figsize=(10.8, 4.5))
    axis.set_xlim(0, 11.7)
    axis.set_ylim(0, 4.8)
    axis.axis("off")
    workflow_band = FancyBboxPatch(
        (3.35, 1.55),
        5.95,
        2.15,
        boxstyle="round,pad=0.12",
        facecolor="#F7FBFB",
        edgecolor="#8AA1AD",
        linewidth=1.0,
        linestyle="--",
        zorder=0,
    )
    axis.add_patch(workflow_band)
    axis.text(
        6.32,
        3.48,
        "LangChain — orquestração do fluxo principal",
        ha="center",
        va="center",
        fontsize=8.5,
        color="#5E6D75",
    )
    boxes = {
        "Frontend\nReact": (0.2, 2.1, 1.25, 1.0, "#E8F5F3"),
        "API\nFastAPI": (1.8, 2.1, 1.25, 1.0, "#E7EEF3"),
        "Modelo ML +\nexplicabilidade": (3.65, 2.1, 1.45, 1.0, "#E7EEF3"),
        "Retriever TF-IDF\n16 sínteses": (5.65, 2.1, 1.45, 1.0, "#FFF5E6"),
        "OpenAI LLM\ncondicional": (7.65, 2.1, 1.45, 1.0, "#F3EAF4"),
        "Análise ao\nprofissional": (9.65, 2.1, 1.55, 1.0, "#E8F5F3"),
        "SQLite\nauditoria": (1.8, 0.3, 1.25, 1.0, "#FBECEF"),
    }
    for label, (x_value, y_value, width, height, color) in boxes.items():
        patch = FancyBboxPatch(
            (x_value, y_value),
            width,
            height,
            boxstyle="round,pad=0.08",
            facecolor=color,
            edgecolor="#48606F",
            linewidth=1.1,
        )
        axis.add_patch(patch)
        axis.text(
            x_value + width / 2,
            y_value + height / 2,
            label,
            ha="center",
            va="center",
            color="#24333D",
            fontweight="bold",
            fontsize=9,
        )

    def arrow(start, end, text=""):
        axis.annotate("", xy=end, xytext=start, arrowprops={"arrowstyle": "->", "color": "#48606F"})
        if text:
            axis.text(
                (start[0] + end[0]) / 2,
                (start[1] + end[1]) / 2 + 0.15,
                text,
                ha="center",
                fontsize=7.5,
                color="#5E6D75",
            )

    arrow((1.45, 2.6), (1.8, 2.6))
    arrow((3.05, 2.6), (3.65, 2.6))
    arrow((5.1, 2.6), (5.65, 2.6))
    arrow((7.1, 2.6), (7.65, 2.6))
    arrow((9.1, 2.6), (9.65, 2.6))
    arrow((2.425, 2.1), (2.425, 1.3))
    axis.set_title("Arquitetura e fluxo principal da Guardiã AI", fontsize=13, pad=12)
    figure.savefig(paths["architecture"], dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(figure)
    return paths


def _cover(document: Document, context: dict[str, Any]) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(50)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("GUARDIÃ AI")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Apoio à triagem de risco materno")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(22)
    run = paragraph.add_run("RELATÓRIO TÉCNICO APROFUNDADO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(INK)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Tech Challenge — Fase 5")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_paragraph().paragraph_format.space_after = Pt(25)
    _add_callout(
        document,
        "DADOS + MACHINE LEARNING + EXPLICABILIDADE + LLM + RAG + APLICAÇÃO",
    )

    today = datetime.now(tz=ZoneInfo("America/Sao_Paulo")).date()
    date_text = f"{today.day} de {MONTHS[today.month]} de {today.year}"
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(25)
    run = paragraph.add_run(
        f"Versão do modelo: 1.0.0\n"
        f"Commit analisado: {context['commit']} ({context['branch']})\n"
        f"Documento gerado em {date_text}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_paragraph().paragraph_format.space_after = Pt(24)
    _add_callout(
        document,
        "AVISO: demonstração acadêmica não validada clinicamente. Não constitui diagnóstico, "
        "prescrição ou decisão automática. A avaliação e a decisão final são humanas.",
        tone="danger",
    )
    document.add_page_break()


def _executive_summary(document: Document, context: dict[str, Any]) -> None:
    _add_heading(document, "Resumo executivo", 1)
    report = context["report"]
    selected = report["evaluations"][report["selected_model"]]
    _add_body(
        document,
        "A Guardiã AI é um protótipo acadêmico de apoio à triagem de risco materno. A aplicação "
        "recebe seis medidas, executa classificação multiclasse, explica a predição, recupera "
        "sínteses de documentos oficiais e, quando uma chave é configurada, solicita a uma LLM "
        "uma explicação contextualizada para o profissional responsável.",
    )
    _add_table(
        document,
        ["Dimensão", "Evidência atual"],
        [
            ["Cenário", "Saúde da mulher — triagem de risco durante a gestação"],
            ["Dataset", "1.014 linhas brutas; 452 após deduplicação exata"],
            ["Modelos", "Regressão logística multinomial e random forest"],
            [
                "Modelo selecionado",
                (
                    f"Random forest — F1 macro {_percent(selected['f1_macro'])}; "
                    f"recall alto {_percent(selected['recall_by_class']['high'])}"
                ),
            ],
            ["RAG", "TF-IDF sobre 16 sínteses de quatro documentos oficiais"],
            ["LLM", "Integração implementada; chamada externa real ainda não comprovada"],
            ["Aplicação", "Frontend React, API FastAPI, auditoria SQLite e Docker Compose"],
        ],
        widths=[4.2, 12.0],
    )
    _add_callout(
        document,
        "Resultado principal: o protótipo cobre o fluxo exigido, mas as métricas têm grande "
        "incerteza, o alvo não é um desfecho clínico auditável e não existe validação para "
        "uso real.",
        tone="warning",
    )
    _add_body(
        document,
        "Este relatório separa três conceitos: aderência ao enunciado acadêmico, evidência "
        "experimental disponível e condições necessárias antes de qualquer uso assistencial.",
    )
    document.add_page_break()
    _add_toc(document)


def _problem_and_scope(document: Document) -> None:
    _add_heading(document, "1. Problema escolhido, atores e escopo", 1)
    _add_heading(document, "1.1 Problema", 2)
    _add_body(
        document,
        "O cenário principal é a triagem inicial de risco materno. Em um atendimento, um "
        "profissional informa idade, pressão arterial, glicemia, temperatura e frequência "
        "cardíaca. O sistema estima a classe aprendida no dataset (baixo, médio ou alto), mostra "
        "probabilidades e fatores contribuintes, recupera informação documental e registra a "
        "análise para consulta posterior.",
    )
    _add_table(
        document,
        ["Ator", "Papel", "Limite"],
        [
            [
                "Pessoa atendida",
                "Origem das medidas e do contexto",
                "Não recebe decisão automática",
            ],
            [
                "Profissional",
                "Interpreta resultado e fontes",
                "Mantém integralmente a decisão final",
            ],
            ["Guardiã AI", "Organiza evidências e produz apoio", "Não diagnostica nem prescreve"],
        ],
        widths=[3.4, 6.2, 6.4],
    )
    _add_heading(document, "1.2 Jornada implementada", 2)
    for step in (
        "Preenchimento das seis variáveis aceitas pelo modelo e de uma pergunta opcional.",
        "Validação do domínio observado no dataset e envio à API.",
        "Predição por Machine Learning e cálculo de probabilidades por classe.",
        "Decomposição local das contribuições para a classe estimada.",
        "Consulta TF-IDF à base de sínteses documentais.",
        "Explicação pela LLM quando configurada, ou fallback explícito quando indisponível.",
        "Exibição de resultado, fontes, aviso de responsabilidade e registro em SQLite.",
    ):
        _add_number(document, step)
    _add_heading(document, "1.3 Dentro e fora do escopo", 2)
    _add_table(
        document,
        ["Dentro do protótipo", "Fora do protótipo"],
        [
            [
                "Classificar o padrão estatístico do dataset",
                "Diagnóstico médico ou prognóstico individual",
            ],
            ["Explicar fatores usados pelo modelo", "Inferir causalidade clínica"],
            ["Recuperar sínteses e links oficiais", "Substituir a leitura do protocolo integral"],
            ["Apoiar perguntas informacionais", "Prescrever exames, medicamentos ou condutas"],
            ["Registrar a análise", "Prontuário eletrônico ou sistema validado para produção"],
        ],
        widths=[8.1, 8.1],
    )


def _data_and_eda(document: Document, context: dict[str, Any], charts: dict[str, Path]) -> None:
    _add_heading(document, "2. Dados utilizados e análise exploratória", 1)
    report = context["report"]
    audit = report["dataset"]["audit"]
    raw = context["raw"]
    clean = context["clean"]
    _add_heading(document, "2.1 Proveniência e integridade", 2)
    _add_body(
        document,
        "Foi utilizado o ML-Ready Maternal Health Risk Assessment Dataset, distribuído no "
        "Kaggle e relacionado ao conjunto Maternal Health Risk do UCI. A descrição atribui a "
        "coleta a hospitais, clínicas comunitárias e centros maternos em Bangladesh. O arquivo "
        "não contém nome ou identificador de paciente, mas também não contém metadados que "
        "permitam auditar a amostragem ou o processo de rotulagem.",
    )
    _add_table(
        document,
        ["Item", "Valor"],
        [
            ["Arquivo", "backend/data/raw/maternal_health_risk.csv"],
            ["SHA-256", report["dataset"]["sha256"]],
            ["Linhas observadas", str(audit["raw_row_count"])],
            ["Licença registrada", "Apache 2.0 no Kaggle; referência UCI sob CC BY 4.0"],
            ["Dados identificáveis", "Nenhum campo identificador no CSV"],
        ],
        widths=[4.2, 12.0],
        font_size=8.0,
    )
    _add_heading(document, "2.2 Dicionário de dados", 2)
    ranges = audit["statistical_outliers_retained"]
    _add_table(
        document,
        ["Campo", "Significado", "Unidade", "Mínimo", "Máximo", "Valores únicos"],
        [
            [
                feature,
                FEATURE_LABELS[feature],
                FEATURE_UNITS[feature],
                _number(ranges[feature]["observed_minimum"], 1),
                _number(ranges[feature]["observed_maximum"], 1),
                int(clean[feature].nunique()),
            ]
            for feature in FEATURE_COLUMNS
        ]
        + [["risk_level", "Classe-alvo 0/1/2", "categoria", "0", "2", "3"]],
        font_size=7.5,
    )
    _add_heading(document, "2.3 Qualidade e distribuição", 2)
    duplicate_rate = audit["exact_duplicate_rows_removed"] / audit["raw_row_count"]
    conflict_vector_rate = (
        audit["feature_vector_groups_with_conflicting_labels"]
        / audit["unique_feature_vector_count"]
    )
    conflict_row_rate = context["conflict_rows"] / audit["unique_exact_row_count"]
    _add_table(
        document,
        ["Achado", "Quantidade", "Interpretação"],
        [
            [
                "Duplicatas exatas",
                f"{audit['exact_duplicate_rows_removed']} ({_percent(duplicate_rate)})",
                "Não é possível distinguir erro de cópia de medições legítimas repetidas.",
            ],
            [
                "Linhas únicas",
                str(audit["unique_exact_row_count"]),
                "Base efetiva pequena para uma classificação de três classes.",
            ],
            [
                "Vetores distintos",
                str(audit["unique_feature_vector_count"]),
                "Atributos fortemente discretizados e repetidos.",
            ],
            [
                "Grupos conflitantes",
                f"35 ({_percent(conflict_vector_rate)} dos vetores)",
                f"Envolvem {context['conflict_rows']} linhas ({_percent(conflict_row_rate)}).",
            ],
            [
                "Ausentes",
                "0",
                "Não há ausentes no arquivo; imputação permanece no pipeline para robustez.",
            ],
            [
                "Extremo HeartRate",
                "7 bpm em duas linhas brutas",
                "Mantido porque não existe evidência auditável para corrigi-lo.",
            ],
        ],
        widths=[3.6, 4.0, 8.6],
        font_size=8.0,
    )
    _add_picture(
        document,
        charts["classes"],
        "Figura 1 — A deduplicação altera materialmente a composição das classes.",
    )
    raw_proportions = raw["risk_level"].value_counts(normalize=True).reindex([0, 1, 2])
    clean_proportions = clean["risk_level"].value_counts(normalize=True).reindex([0, 1, 2])
    _add_body(
        document,
        "Antes da deduplicação, baixo/médio/alto representam "
        f"{_percent(raw_proportions[0])}, {_percent(raw_proportions[1])} e "
        f"{_percent(raw_proportions[2])}. Depois, passam a "
        f"{_percent(clean_proportions[0])}, {_percent(clean_proportions[1])} e "
        f"{_percent(clean_proportions[2])}. A política de duplicatas é, portanto, uma decisão "
        "metodológica relevante e não uma limpeza neutra.",
    )
    _add_picture(
        document,
        charts["conflicts"],
        "Figura 2 — Tipos de conflito encontrados entre vetores idênticos.",
    )
    _add_body(
        document,
        "Os conflitos incluem 24 grupos baixo–médio, cinco baixo–alto, cinco médio–alto e um "
        "grupo associado às três classes. Isso indica ruído de rótulo ou perda de variáveis "
        "clínicas relevantes no dataset.",
    )
    _add_picture(
        document,
        charts["distributions"],
        "Figura 3 — Distribuições univariadas após a remoção de duplicatas exatas.",
    )
    statistics = audit["descriptive_statistics_after_exact_deduplication"]
    _add_table(
        document,
        ["Atributo", "Média", "Desvio", "Mín.", "Q1", "Mediana", "Q3", "Máx."],
        [
            [
                FEATURE_LABELS[feature],
                _number(statistics[feature]["mean"]),
                _number(statistics[feature]["std"]),
                _number(statistics[feature]["min"], 1),
                _number(statistics[feature]["25%"], 1),
                _number(statistics[feature]["50%"], 1),
                _number(statistics[feature]["75%"], 1),
                _number(statistics[feature]["max"], 1),
            ]
            for feature in FEATURE_COLUMNS
        ],
        font_size=7.2,
    )
    _add_picture(
        document,
        charts["correlation"],
        "Figura 4 — Correlações exploratórias; classe codificada ordinalmente apenas para "
        "inspeção.",
    )
    _add_callout(
        document,
        "A correlação com o código 0/1/2 não prova relação clínica nem causalidade. O alvo não é "
        "um desfecho prospectivo e seu método de rotulagem não é auditável pelo CSV.",
        tone="warning",
    )


def _preparation(document: Document, context: dict[str, Any]) -> None:
    _add_heading(document, "3. Preparação dos dados e prevenção de vazamento", 1)
    report = context["report"]
    split = report["split"]
    _add_heading(document, "3.1 Pipeline", 2)
    for step in (
        "Verificação do SHA-256 e leitura imutável do CSV incorporado.",
        "Normalização dos nomes das colunas e seleção dos seis atributos e do alvo.",
        "Conversão numérica; infinitos são tratados como ausentes.",
        "Mapeamento do alvo para 0=baixo, 1=médio e 2=alto.",
        "Remoção de repetições exatas da linha completa.",
        "Agrupamento por vetor completo de atributos, inclusive nos casos de rótulo conflitante.",
        "StratifiedGroupKFold para impedir que um vetor idêntico apareça no treino e no teste.",
        "Imputação ajustada somente na partição de treino; padronização apenas na regressão.",
    ):
        _add_number(document, step)
    _add_heading(document, "3.2 Partição de avaliação", 2)
    _add_table(
        document,
        ["Partição", "Linhas", "Baixo", "Médio", "Alto", "Grupos"],
        [
            [
                "Treino",
                split["train_row_count"],
                split["train_class_counts"]["low"],
                split["train_class_counts"]["mid"],
                split["train_class_counts"]["high"],
                split["train_feature_group_count"],
            ],
            [
                "Holdout",
                split["test_row_count"],
                split["test_class_counts"]["low"],
                split["test_class_counts"]["mid"],
                split["test_class_counts"]["high"],
                split["test_feature_group_count"],
            ],
        ],
    )
    _add_callout(
        document,
        f"Sobreposição de grupos entre treino e holdout: {split['feature_group_overlap_count']}. "
        "Esse controle evita vazamento por vetores idênticos, mas não corrige ruído de rótulo.",
    )
    _add_body(
        document,
        "A deduplicação foi necessária para reduzir uma fonte evidente de otimismo, mas pode "
        "remover observações legítimas. Como não existem paciente, data ou sequência temporal, "
        "não há como distinguir repetição indevida de recorrência real. Uma análise de "
        "sensibilidade com políticas alternativas ainda não foi realizada.",
    )


def _modeling(document: Document, context: dict[str, Any], charts: dict[str, Path]) -> None:
    _add_heading(document, "4. Modelos, protocolo experimental e métricas", 1)
    report = context["report"]
    evaluations = report["evaluations"]
    _add_heading(document, "4.1 Modelos testados", 2)
    _add_table(
        document,
        ["Modelo", "Pré-processamento", "Hiperparâmetros relevantes"],
        [
            [
                "Regressão logística multinomial",
                "Imputação mediana + StandardScaler",
                "class_weight=balanced; solver=lbfgs; max_iter=3000",
            ],
            [
                "Random forest",
                "Imputação mediana",
                "300 árvores; min_samples_leaf=2; class_weight=balanced; seed=42",
            ],
        ],
        widths=[4.6, 5.1, 6.6],
    )
    _add_body(
        document,
        "A regra de seleção foi definida para maximizar primeiro o recall de alto risco, depois "
        "F1 macro e, por fim, accuracy. A escolha prioriza reduzir registros altos classificados "
        "como outra classe, mas não atribui significado clínico ao rótulo do dataset.",
    )
    _add_heading(document, "4.2 Comparação agregada", 2)
    _add_table(
        document,
        [
            "Modelo",
            "Accuracy",
            "Precision macro",
            "Recall macro",
            "F1 macro",
            "Recall alto",
            "FN alto",
        ],
        [
            [
                MODEL_LABELS[name],
                _percent(evaluations[name]["accuracy"]),
                _percent(evaluations[name]["precision_macro"]),
                _percent(evaluations[name]["recall_macro"]),
                _percent(evaluations[name]["f1_macro"]),
                _percent(evaluations[name]["recall_by_class"]["high"]),
                evaluations[name]["high_risk_false_negatives"],
            ]
            for name in report["candidate_models"]
        ],
        font_size=7.2,
    )
    _add_picture(
        document,
        charts["confusion"],
        "Figura 5 — Matrizes de confusão dos dois candidatos no mesmo holdout agrupado.",
    )
    _add_heading(document, "4.3 Resultados por classe", 2)
    for name in report["candidate_models"]:
        _add_heading(document, MODEL_LABELS[name], 3)
        class_metrics = _per_class_metrics(evaluations[name]["confusion_matrix"])
        _add_table(
            document,
            ["Classe", "Precision", "Recall", "F1", "Suporte"],
            [
                [
                    item["class"],
                    _percent(float(item["precision"])),
                    _percent(float(item["recall"])),
                    _percent(float(item["f1"])),
                    item["support"],
                ]
                for item in class_metrics
            ],
        )
    high_interval = _wilson_interval(20, 22)
    accuracy_interval = _wilson_interval(62, 89)
    _add_callout(
        document,
        "A random forest obteve 20 acertos entre 22 registros altos: recall 90,91%, mas IC "
        f"Wilson 95% de {_percent(high_interval[0])} a {_percent(high_interval[1])}. A accuracy "
        f"de 62/89 tem IC aproximado de {_percent(accuracy_interval[0])} a "
        f"{_percent(accuracy_interval[1])}. A amostra não sustenta precisão excessiva.",
        tone="warning",
    )
    _add_heading(document, "4.4 Impacto dos erros", 2)
    _add_table(
        document,
        ["Erro observado na random forest", "Quantidade", "Impacto potencial no apoio"],
        [
            ["Alto → médio", "2", "Pode reduzir a percepção de prioridade do caso."],
            [
                "Médio → baixo",
                "9",
                "É o maior erro de rebaixamento e evidencia fragilidade na classe média.",
            ],
            ["Baixo → médio", "10", "Pode aumentar revisões e carga profissional desnecessária."],
            [
                "Baixo → alto",
                "1",
                "Pode gerar alarme excessivo, embora a decisão permaneça humana.",
            ],
            [
                "Médio → alto",
                "5",
                "Pode elevar casos ambíguos e reduzir a precisão da classe alta.",
            ],
        ],
        widths=[4.2, 2.3, 9.8],
    )
    _add_body(
        document,
        "A classe média é o principal ponto fraco: somente 7 de 21 registros foram corretamente "
        "classificados, com precision 36,84%, recall 33,33% e F1 35,00%. Um sistema real deveria "
        "tratar essa incerteza explicitamente, em vez de apresentar apenas a classe vencedora.",
    )
    _add_heading(document, "4.5 Limitações do desenho experimental", 2)
    for item in (
        (
            "O mesmo holdout de 89 linhas compara os candidatos e fundamenta a escolha do "
            "vencedor; não existe teste final independente após a seleção."
        ),
        (
            "O estimador publicado é reajustado nas 452 linhas limpas; suas métricas continuam "
            "sendo as do modelo anterior treinado somente na partição de treino."
        ),
        (
            "Não houve busca sistemática de hiperparâmetros, validação cruzada repetida ou "
            "validação externa."
        ),
        (
            "As probabilidades exibidas não foram avaliadas por calibração, Brier score ou "
            "reliability curve."
        ),
        (
            "Não existem análise por subgrupo, avaliação formal de viés ou intervalo de "
            "confiança no artefato original."
        ),
    ):
        _add_bullet(document, item)


def _explainability(document: Document, context: dict[str, Any], charts: dict[str, Path]) -> None:
    _add_heading(document, "5. Explicabilidade do modelo", 1)
    report = context["report"]
    explanation = context["explanation"]
    prediction = context["prediction"]
    _add_heading(document, "5.1 Importância global", 2)
    _add_body(
        document,
        "A random forest usa importância por redução de impureza. Ela resume quanto cada atributo "
        "participou das divisões das árvores, mas pode favorecer variáveis com maior variedade de "
        "pontos de corte. Não representa causalidade nem relevância clínica independente.",
    )
    _add_picture(
        document,
        charts["importance"],
        "Figura 6 — Importância global normalizada da random forest final.",
    )
    _add_table(
        document,
        ["Atributo", "Importância"],
        [
            [FEATURE_LABELS[feature], _percent(value)]
            for feature, value in report["global_feature_importance"].items()
        ],
        widths=[8.0, 4.0],
    )
    _add_heading(document, "5.2 Explicação local reproduzível", 2)
    features = context["example_features"]
    _add_table(
        document,
        ["Entrada", "Valor"],
        [
            [FEATURE_LABELS[feature], f"{features[feature]:g} {FEATURE_UNITS[feature]}"]
            for feature in FEATURE_COLUMNS
        ],
        widths=[8.0, 4.0],
    )
    _add_table(
        document,
        ["Classe", "Probabilidade"],
        [
            ["Baixo", _percent(prediction["probabilities"]["low"])],
            ["Médio", _percent(prediction["probabilities"]["mid"])],
            ["Alto", _percent(prediction["probabilities"]["high"])],
        ],
        widths=[8.0, 4.0],
    )
    _add_body(
        document,
        "A explicação percorre o caminho da observação em cada árvore. A diferença de "
        "probabilidade "
        "entre pai e filho é atribuída ao atributo usado na divisão; a média das árvores produz "
        "uma decomposição exata da probabilidade da classe explicada.",
    )
    _add_picture(
        document,
        charts["waterfall"],
        "Figura 7 — Baseline e contribuições assinadas para a classe alta no estudo de caso.",
    )
    contribution_rows = []
    for feature, value in sorted(
        explanation["feature_contributions"].items(),
        key=lambda item: abs(item[1]),
        reverse=True,
    ):
        contribution_rows.append(
            [
                FEATURE_LABELS[feature],
                f"{features[feature]:g} {FEATURE_UNITS[feature]}",
                f"{value * 100:+.3f} pp".replace(".", ","),
                "Aumenta" if value > 0 else "Reduz" if value < 0 else "Neutra",
            ]
        )
    _add_table(
        document,
        ["Atributo", "Valor", "Contribuição", "Direção"],
        contribution_rows,
        widths=[5.2, 3.5, 3.7, 3.4],
    )
    _add_callout(
        document,
        f"Baseline alto: {_percent(explanation['baseline_probability'], 3)}. Probabilidade "
        f"reconstruída: {_percent(explanation['reconstructed_probability'], 3)}. Erro numérico: "
        f"{explanation['reconstruction_error']:.2e}. A igualdade verifica o cálculo, não a "
        "validade clínica da predição.",
    )
    _add_heading(document, "5.3 Limitações da explicação", 2)
    for item in (
        "A explicação descreve o comportamento do estimador, não mecanismos causais.",
        (
            "A interface retorna magnitude, direção e erro, mas não expõe o baseline nem toda a "
            "reconstrução mostrada neste estudo de caso."
        ),
        (
            "Somente a classe estimada é explicada; não há contraste formal entre classes "
            "concorrentes."
        ),
        (
            "Não foi feita comparação com permutation importance, SHAP ou teste de estabilidade "
            "das explicações."
        ),
    ):
        _add_bullet(document, item)


def _architecture(document: Document, charts: dict[str, Path]) -> None:
    _add_heading(document, "6. Arquitetura da aplicação", 1)
    _add_picture(
        document,
        charts["architecture"],
        "Figura 8 — Frontend e backend separados; LangChain organiza o fluxo principal.",
        width=6.7,
    )
    _add_table(
        document,
        ["Componente", "Tecnologia", "Responsabilidade"],
        [
            [
                "Frontend",
                "React + TypeScript + Vite",
                "Formulário, resultado, histórico e métricas",
            ],
            ["Proxy", "nginx", "Entrega da SPA e encaminhamento de /api"],
            ["API", "FastAPI + Pydantic", "Contrato, validação e composição do caso de uso"],
            ["Orquestração", "LangChain RunnableLambda", "ML → recuperação → explicação"],
            ["Modelo", "scikit-learn/joblib", "Predição, probabilidades e explicabilidade"],
            ["RAG", "TF-IDF + cosseno", "Recuperação lexical local"],
            ["LLM", "ChatOpenAI", "Explicação fundamentada quando configurada"],
            ["Auditoria", "SQLite", "Requisição e resposta completas"],
            ["Execução", "Dockerfiles + Compose", "Ambiente reproduzível"],
        ],
        widths=[3.0, 4.7, 8.5],
        font_size=8.0,
    )
    _add_heading(document, "6.1 Contrato HTTP", 2)
    _add_table(
        document,
        ["Método", "Rota", "Finalidade"],
        [
            ["POST", "/api/v1/analyses", "Executar a jornada e persistir o resultado"],
            ["GET", "/api/v1/analyses", "Listar análises recentes"],
            ["GET", "/api/v1/analyses/{id}", "Recuperar análise completa"],
            ["GET", "/api/v1/model/metrics", "Expor auditoria e métricas do treinamento"],
            ["GET", "/health", "Informar modelo, base RAG e estado da LLM"],
        ],
        widths=[2.2, 5.2, 8.8],
    )


def _llm(document: Document) -> None:
    _add_heading(document, "7. Integração com a Large Language Model", 1)
    _add_heading(document, "7.1 Implementação", 2)
    _add_body(
        document,
        "A integração usa ChatOpenAI dentro de uma cadeia LangChain. O modelo padrão configurável "
        "é gpt-5.6-luna e a chave é lida de variável de ambiente. A ausência de chave não é "
        "mascarada: a API retorna llm_used=false e informa que somente ML e recuperação foram "
        "executados.",
    )
    _add_table(
        document,
        ["Informação enviada quando habilitada", "Finalidade"],
        [
            ["Classe e probabilidades", "Explicar o resultado estatístico"],
            ["Seis valores e contribuições", "Contextualizar os fatores do modelo"],
            ["Pergunta livre", "Responder à necessidade informacional do profissional"],
            ["Quatro sínteses recuperadas", "Restringir a resposta ao contexto RAG"],
            ["Título, referência e URL", "Permitir citações [1], [2] e rastreabilidade"],
        ],
        widths=[6.7, 9.5],
    )
    _add_heading(document, "7.2 Restrições do prompt", 2)
    _add_table(
        document,
        ["Regra", "Objetivo"],
        [
            ["Português claro e objetivo", "Adequar a explicação ao profissional"],
            ["Nunca diagnosticar ou prescrever", "Manter o limite de apoio"],
            ["Usar somente predição e sínteses", "Reduzir afirmações sem suporte"],
            ["Declarar insuficiência", "Evitar certeza sem dados"],
            ["Diferenciar ML de documento", "Não atribuir a classificação ao protocolo"],
            ["Citar [1], [2]", "Relacionar resposta às fontes recuperadas"],
        ],
        widths=[6.7, 9.5],
    )
    _add_callout(
        document,
        "Estado da evidência: a integração está implementada, mas a suíte automatizada usa "
        "openai_api_key=None e valida o fallback. A resposta com llm_used=true no frontend é "
        "simulada. Não foi registrada uma chamada real que possa ser apresentada como resultado.",
        tone="warning",
    )
    _add_heading(document, "7.3 Riscos residuais", 2)
    for item in (
        "Citações e groundedness dependem do prompt; não existe validador pós-geração.",
        (
            "Não há teste de prompt injection, contradição, pergunta fora do escopo ou citação "
            "inexistente."
        ),
        (
            "Temperatura, limite de saída e versão do prompt não são persistidos como metadados "
            "auditáveis."
        ),
        (
            "O campo livre pode conter dado identificável e ser enviado ao provedor apesar do "
            "aviso da interface."
        ),
        "A chamada síncrona pode bloquear a rota durante timeout e retries.",
    ):
        _add_bullet(document, item)


def _rag(document: Document, context: dict[str, Any], charts: dict[str, Path]) -> None:
    _add_heading(document, "8. Funcionamento do RAG", 1)
    _add_heading(document, "8.1 Base de conhecimento", 2)
    _add_body(
        document,
        "A base contém 16 sínteses próprias e rastreáveis de quatro documentos oficiais. Os PDFs "
        "integrais não são redistribuídos; cada item guarda título, URL e referência de página ou "
        "seção. As sínteses não substituem a fonte integral.",
    )
    _add_table(
        document,
        ["Documento", "Ano", "Seções", "Licença registrada", "Emissor"],
        [
            [
                source["title"],
                source["year"],
                source["sections"],
                source["license"],
                source["issuer"],
            ]
            for source in context["source_metadata"]
        ],
        font_size=7.0,
    )
    _add_heading(document, "8.2 Recuperação", 2)
    for step in (
        "Normalização da pergunta, classe estimada e três fatores mais relevantes.",
        "Vetorização TF-IDF somente do texto das 16 sínteses, com unigramas e bigramas.",
        "Similaridade de cosseno entre consulta e cada síntese.",
        "Ordenação decrescente e retorno das quatro primeiras seções.",
        "Inclusão de título, referência, URL, síntese e score na resposta auditada.",
    ):
        _add_number(document, step)
    _add_body(document, f"Consulta reproduzida: “{context['retrieval_query']}”")
    _add_table(
        document,
        ["Posição", "Documento", "Referência", "Score"],
        [
            [index, item.title, item.reference, f"{item.relevance_score:.6f}"]
            for index, item in enumerate(context["retrieval"], 1)
        ],
        widths=[1.6, 5.4, 7.3, 2.0],
        font_size=7.3,
    )
    _add_picture(
        document,
        charts["retrieval"],
        "Figura 9 — Scores reais do retriever para a consulta do estudo de caso.",
    )
    _add_heading(document, "8.3 Limitações e avaliação", 2)
    for item in (
        "A ingestão não extrai automaticamente os PDFs; o corpus é composto por sínteses manuais.",
        "Título, emissor e referência não participam da vetorização, apenas section.text.",
        (
            "O retriever sempre devolve top_k, inclusive quando a similaridade é zero; não há "
            "limiar de relevância."
        ),
        "Há somente dois testes temáticos, para pressão e glicemia; não existe conjunto-ouro.",
        "Precision@K, Recall@K, MRR, cobertura e groundedness ainda não foram medidos.",
        "Uma fonte recuperada não garante que a LLM a citará corretamente.",
    ):
        _add_bullet(document, item)


def _application(document: Document) -> None:
    _add_heading(document, "9. Aplicação e jornada do profissional", 1)
    _add_heading(document, "9.1 Interface", 2)
    _add_table(
        document,
        ["Área", "Conteúdo"],
        [
            ["Nova análise", "Seis variáveis, pergunta opcional, faixas e aviso de apoio"],
            ["Resultado", "Classe, probabilidades, explicação, contribuições e fontes"],
            ["Rastreabilidade", "UUID, data, modelo, versão e dados utilizados"],
            ["Histórico", "Vinte registros recentes e expansão dos dados de entrada"],
            ["Desempenho", "Métricas e matriz publicadas pelo backend"],
        ],
        widths=[4.2, 12.0],
    )
    _add_heading(document, "9.2 Pontos fortes", 2)
    for item in (
        "Frontend e backend são independentes e conectados por contrato HTTP.",
        "A interface distingue classificação estimada de decisão profissional.",
        "Fontes recuperadas permanecem visíveis mesmo sem execução da LLM.",
        "Entradas, resultado e versão do modelo podem ser consultados posteriormente.",
        "Estados de carregamento, erro e LLM indisponível são tratados.",
    ):
        _add_bullet(document, item)
    _add_heading(document, "9.3 Limitações da jornada", 2)
    for item in (
        (
            "Temperatura exige Fahrenheit e glicemia mmol/L, unidades pouco usuais no "
            "atendimento brasileiro, sem conversor na interface."
        ),
        (
            "HeartRate=7 é aceito porque pertence ao domínio bruto, apesar de ser extremo não "
            "verificável."
        ),
        (
            "O histórico não reexibe a explicação, as probabilidades e as fontes completas na "
            "própria tabela."
        ),
        "O score de relevância existe na API, mas não é mostrado ao usuário.",
        "Não existe registro explícito da revisão ou decisão final do profissional.",
        "Não existem busca, exportação, retenção ou exclusão de análises.",
    ):
        _add_bullet(document, item)


def _security_and_audit(document: Document) -> None:
    _add_heading(document, "10. Segurança, responsabilidade e auditoria", 1)
    _add_heading(document, "10.1 Controles implementados", 2)
    _add_table(
        document,
        ["Controle", "Implementação"],
        [
            ["Supervisão humana", "Disclaimers no frontend, API, prompt e relatório"],
            ["Limites da LLM", "Proibição de diagnóstico, prescrição e decisão automática"],
            ["Dados estruturados", "Nenhum campo de nome, documento ou prontuário"],
            ["Validação", "Pydantic recusa valores fora do domínio observado"],
            ["Auditoria", "UUID/UTC e persistência da requisição e resposta completas"],
            ["Fallback", "Falha ou ausência da LLM é informada sem texto simulado"],
            ["Segredos", "Chave OpenAI somente em variável de ambiente"],
        ],
        widths=[4.5, 11.7],
    )
    _add_heading(document, "10.2 Conteúdo auditado", 2)
    _add_body(
        document,
        "O SQLite registra classe, rótulo, modelo, versão, uso da LLM, request_json e "
        "response_json. A resposta inclui entrada, probabilidades, contribuições, método, erro de "
        "reconstrução, explicação, fontes e disclaimer. Isso cumpre o requisito acadêmico de saber "
        "quais informações foram usadas e qual resultado foi produzido.",
    )
    _add_heading(document, "10.3 Ameaças e risco residual", 2)
    _add_table(
        document,
        ["Ameaça", "Controle atual", "Risco residual / ação necessária"],
        [
            [
                "Acesso ao histórico",
                "Execução local e UUID",
                "Sem autenticação/autorização; não expor à internet.",
            ],
            [
                "Dados em repouso",
                "Volume Docker persistente",
                "SQLite e JSON em claro; definir criptografia e retenção.",
            ],
            [
                "PII na pergunta",
                "Aviso para não inserir identificadores",
                "Sem detecção/redação antes de persistir e enviar.",
            ],
            [
                "Prompt injection",
                "Mensagem de sistema restritiva",
                "Sem classificação de entrada ou teste adversarial.",
            ],
            [
                "Alucinação/citação falsa",
                "RAG e instrução de citar",
                "Sem validação de citação ou groundedness.",
            ],
            [
                "Alteração do log",
                "Banco local",
                "Sem trilha imutável, assinatura ou segregação de acesso.",
            ],
            [
                "Uso clínico indevido",
                "Avisos de não validação",
                "Exige governança, validação e treinamento dos usuários.",
            ],
        ],
        widths=[3.8, 4.9, 7.5],
        font_size=7.5,
    )
    _add_callout(
        document,
        "O protótipo não demonstra conformidade LGPD nem segurança para produção. Deve usar "
        "somente dados fictícios durante a avaliação acadêmica.",
        tone="danger",
    )


def _testing(document: Document, context: dict[str, Any]) -> None:
    _add_heading(document, "11. Testes, execução e reprodutibilidade", 1)
    _add_heading(document, "11.1 Evidência automatizada", 2)
    _add_table(
        document,
        ["Verificação", "Resultado observado", "Cobertura funcional"],
        [
            ["Backend pytest", "14/14 aprovados", "ML, RAG, API e SQLite"],
            ["Frontend Vitest", "6/6 aprovados", "Formulário, contrato, histórico e métricas"],
            ["Ruff", "Aprovado", "Estilo e erros estáticos Python"],
            ["Vite build", "Aprovado", "TypeScript e bundle de produção"],
            ["Docker", "Imagens e Compose validados", "SPA, proxy, API e healthcheck"],
            ["LLM externa", "Não executada", "Somente fallback real e resposta frontend simulada"],
        ],
        widths=[4.1, 4.5, 7.6],
    )
    _add_heading(document, "11.2 Execução", 2)
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.add_run(
        "cp .env.example .env\n"
        "# preencher OPENAI_API_KEY somente para testar a LLM\n"
        "docker compose up --build"
    ).font.name = "Consolas"
    _add_body(
        document,
        "Aplicação: http://localhost:8080 | API: http://localhost:8000 | "
        "OpenAPI: http://localhost:8000/docs",
    )
    _add_heading(document, "11.3 Treinamento", 2)
    paragraph = document.add_paragraph()
    paragraph.add_run(
        "cd backend\n"
        "python -m training.train \\\n"
        "  --data data/raw/maternal_health_risk.csv \\\n"
        "  --output-dir artifacts \\\n"
        "  --model-version 1.0.0"
    ).font.name = "Consolas"
    _add_table(
        document,
        ["Elemento", "Registro"],
        [
            ["Semente", str(context["report"]["random_state"])],
            ["Dataset SHA-256", context["report"]["dataset"]["sha256"]],
            ["Versão do modelo", context["report"]["model_version"]],
            ["Artefato", "maternal_risk_model_v1.0.0.joblib"],
            ["Métricas", "training_report_v1.0.0.json"],
            ["Commit analisado", context["commit"]],
        ],
        widths=[4.2, 12.0],
        font_size=8.0,
    )


def _limitations(document: Document) -> None:
    _add_heading(document, "12. Principais limitações", 1)
    categories = {
        "Dados": [
            "Somente 452 linhas após deduplicação; 55,42% do arquivo bruto foi removido.",
            (
                "Quinze vírgula setenta e um por cento das linhas limpas pertencem a grupos com "
                "rótulos conflitantes."
            ),
            "Sem paciente, tempo, idade gestacional, sintomas, histórico ou desfecho clínico.",
            "Origem atribuída a Bangladesh e processo de rotulagem não auditável.",
        ],
        "Modelagem": [
            "Um único holdout pequeno foi usado para comparar e selecionar o modelo.",
            "Sem teste final independente, validação externa, calibração ou análise por subgrupo.",
            "Recall médio de 33,33% e intervalo amplo para recall alto.",
            "Probabilidades são saídas do estimador, não riscos clínicos calibrados.",
        ],
        "Explicabilidade": [
            "Importância de impureza pode ser enviesada e não foi validada por outro método.",
            "Contribuições explicam o estimador, não causalidade ou prognóstico.",
        ],
        "RAG e LLM": [
            "Corpus de apenas 16 sínteses manuais e recuperação lexical sem avaliação formal.",
            "Ausência de limiar de relevância e de validador de citações/groundedness.",
            "Nenhuma chamada externa real está documentada como evidência.",
        ],
        "Aplicação e segurança": [
            "Sem autenticação, autorização, criptografia aplicada, retenção ou exclusão.",
            "Pergunta livre e valores clínicos são persistidos e podem ser enviados ao provedor.",
            "Unidades de entrada não são adaptadas ao fluxo brasileiro.",
            "Não existe validação clínica, ética, prospectiva ou estudo com usuários.",
        ],
    }
    for category, items in categories.items():
        _add_heading(document, f"12.{list(categories).index(category) + 1} {category}", 2)
        for item in items:
            _add_bullet(document, item)
    _add_callout(
        document,
        "Consequência: a Guardiã AI demonstra integração tecnológica acadêmica, mas não deve ser "
        "usada em atendimento real sem redesenho metodológico, segurança, governança e validação.",
        tone="danger",
    )


def _conclusion(document: Document) -> None:
    _add_heading(document, "13. Conclusão e evolução recomendada", 1)
    _add_body(
        document,
        "A solução transforma um classificador isolado em uma aplicação funcional com interface, "
        "explicabilidade, RAG, integração LLM e auditoria. Isso atende ao núcleo técnico do Tech "
        "Challenge. A contribuição mais sólida é a prevenção de vazamento entre vetores idênticos "
        "e a decomposição local exata da random forest.",
    )
    _add_body(
        document,
        "A evidência disponível, entretanto, não permite afirmar confiabilidade clínica. O "
        "dataset é pequeno, ruidoso e sem desfecho; o experimento não possui teste independente; "
        "RAG e LLM não foram avaliados sistematicamente; e a aplicação não possui controles de "
        "produção. A apresentação deve assumir esses limites com clareza.",
    )
    _add_heading(document, "13.1 Próximos passos priorizados", 2)
    priorities = [
        ("P0", "Versionar e publicar o README obrigatório no repositório Git."),
        ("P0", "Executar uma chamada real da LLM, guardar evidência e verificar as citações."),
        ("P0", "Gravar e publicar o vídeo de até 15 minutos com a jornada completa."),
        ("P1", "Separar seleção e teste final ou adotar validação cruzada agrupada repetida."),
        ("P1", "Avaliar calibração, intervalos de confiança e desempenho por subgrupo."),
        ("P1", "Criar conjunto-ouro para RAG e medir Precision@K, Recall@K e groundedness."),
        ("P1", "Adicionar limiar de relevância, validação de citações e testes adversariais."),
        ("P2", "Rever unidades, registrar decisão humana e melhorar a consulta do histórico."),
        (
            "P2",
            "Definir autenticação, criptografia, retenção e anonimização antes de qualquer piloto.",
        ),
    ]
    _add_table(
        document, ["Prioridade", "Ação"], [list(item) for item in priorities], widths=[2.4, 13.8]
    )


def _traceability(document: Document, context: dict[str, Any]) -> None:
    _add_heading(document, "Apêndice A — Matriz de rastreabilidade do enunciado", 1)
    readme_status = (
        "Atendido no snapshot Git"
        if context["readme_tracked"]
        else "Pendente: README local não rastreado"
    )
    rows = [
        ["Cenário de saúde da mulher", "Relatório, API e frontend", "Atendido"],
        ["Dataset relacionado", "CSV materno + licença + SHA-256", "Atendido"],
        ["Explorar e preparar dados", "train.py e seções 2–3", "Atendido"],
        ["Pelo menos dois modelos", "Logística e random forest", "Atendido"],
        ["Métricas e impactos dos erros", "JSON, matrizes e seção 4", "Atendido"],
        ["Interpretação da predição", "Importância + decomposição local", "Atendido"],
        ["Integração LLM", "ChatOpenAI/LangChain", "Implementada; execução real pendente"],
        ["RAG", "TF-IDF, 16 sínteses, fontes", "Atendido no mínimo"],
        ["LangChain ou LangGraph", "RunnableLambda no fluxo principal", "Atendido"],
        ["Decisão humana e limites", "Prompt, API, UI e relatório", "Atendido"],
        ["Registro das análises", "SQLite com request/response", "Atendido"],
        ["Código e treinamento em Git", f"Commit {context['commit']}", "Atendido"],
        ["README e instruções em Git", "README.md existe localmente", readme_status],
        ["Interface/API", "React + FastAPI", "Atendido"],
        ["Dockerfile", "Dockerfiles + Compose", "Atendido"],
        ["Relatório técnico PDF", "docs/relatorio-tecnico.pdf", "Atendido"],
        ["Relatório técnico DOCX", "docs/relatorio-tecnico.docx", "Atendido após geração"],
        ["Vídeo ≤15 min YouTube/Vimeo", "Apenas roteiro disponível", "Faltando"],
    ]
    _add_table(
        document,
        ["Requisito", "Evidência", "Status"],
        rows,
        widths=[6.0, 6.2, 4.1],
        font_size=7.5,
    )


def _references(document: Document, context: dict[str, Any]) -> None:
    _add_heading(document, "Referências", 1)
    references = [
        (
            "Dataset Kaggle",
            (
                "https://www.kaggle.com/datasets/arshmankhalid/"
                "ml-ready-maternal-health-risk-assessment-dataset"
            ),
        ),
        ("UCI Maternal Health Risk / DOI", "https://doi.org/10.24432/C5DP5D"),
    ]
    references.extend((source["title"], source["url"]) for source in context["source_metadata"])
    references.append(("OpenAI — modelos", "https://developers.openai.com/api/docs/models"))
    for index, (title, url) in enumerate(references, 1):
        _add_body(document, f"{index}. {title}. {url}")
    _add_heading(document, "Artefatos internos consultados", 2)
    for item in (
        "backend/data/raw/maternal_health_risk.csv",
        "backend/artifacts/training_report_v1.0.0.json",
        "backend/artifacts/maternal_risk_model_v1.0.0.joblib",
        "backend/training/train.py",
        "backend/app/services/workflow.py",
        "backend/app/rag/retriever.py",
        "backend/app/rag/explainer.py",
        "docs/relatorio-tecnico.pdf",
    ):
        _add_bullet(document, item)


def build_document(context: dict[str, Any], charts: dict[str, Path]) -> Document:
    document = Document()
    _configure_document(document)
    properties = document.core_properties
    properties.title = "Guardiã AI — Relatório técnico aprofundado"
    properties.subject = "Tech Challenge Fase 5 — Triagem de risco materno"
    properties.author = "Equipe Guardiã AI"
    properties.keywords = "Machine Learning, LLM, RAG, saúde materna, explicabilidade"
    properties.comments = "Gerado por docs/generate_report_docx.py com artefatos auditáveis."

    _cover(document, context)
    _executive_summary(document, context)
    _problem_and_scope(document)
    _data_and_eda(document, context, charts)
    _preparation(document, context)
    _modeling(document, context, charts)
    _explainability(document, context, charts)
    _architecture(document, charts)
    _llm(document)
    _rag(document, context, charts)
    _application(document)
    _security_and_audit(document)
    _testing(document, context)
    _limitations(document)
    _conclusion(document)
    _traceability(document, context)
    _references(document, context)
    return document


def main() -> int:
    context = _load_context()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="guardia-ai-docx-") as temporary:
        charts = _create_charts(context, Path(temporary))
        document = build_document(context, charts)
        document.save(OUTPUT)
    print(f"Relatório DOCX gerado: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
